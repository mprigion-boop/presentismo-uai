import streamlit as st
import pandas as pd
import os
import io
from docx import Document

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(
    page_title="Control de Asistencia - Premium", 
    layout="wide",
    page_icon="🤖"
)

# --- CABECERA DE LA APP ---
col_logo, col_titulo = st.columns([1, 5])

with col_logo:
    # Recuerda descomentar esta línea y poner tu logo cuando estés listo
    # st.image("TU_ARCHIVO_DE_LOGO.png", width=120) 
    st.markdown("**(Logo aquí)**")

with col_titulo:
    st.title("Control de Asistencia - UAIONLINE")

# 1. Subida del archivo
uploaded_file = st.sidebar.file_uploader("Sube el archivo CSV del campus", type=["csv"])

if uploaded_file is not None:
    file_name_no_ext = os.path.splitext(uploaded_file.name)[0]
    
    # Lectura y limpieza inicial
    df = pd.read_csv(uploaded_file, sep=',', encoding='utf-8-sig')
    df.columns = df.columns.str.strip()

    # Columnas clave
    col_dias = 'Días desde el último acceso'
    col_nombre_completo = 'Nombre Completo'
    
    # Validación de seguridad
    if 'Apellido' in df.columns and 'Nombre' in df.columns:
        df[col_nombre_completo] = df['Apellido'] + ", " + df['Nombre']
    else:
        st.error("Error: No se encontraron las columnas 'Apellido' o 'Nombre'.")
        st.stop()

    st.subheader(f"Reporte: {file_name_no_ext}")
    
    # Creamos las pestañas
    tab1, tab2, tab3 = st.tabs(["📊 Panel de Control", "👥 Listado de Alumnos", "📥 Centro de Exportación"])

    # ==========================================
    # PESTAÑA 1: PANEL DE CONTROL
    # ==========================================
    with tab1:
        st.markdown("#### Configuración de Alertas")
        
        dias_limite = st.slider(
            "Definir límite de días para considerar Ausente:", 
            min_value=1, max_value=60, value=15
        )
        st.divider()

        # 1. Lógica de Asistencia
        def clasificar_dinamico(dias):
            if pd.isna(dias) or dias >= dias_limite:
                return "❌ AUSENTE"
            return "✅ PRESENTE"

        df['Estado'] = df[col_dias].apply(clasificar_dinamico)

        # 2. Lógica de Limpieza de Días (NUEVO BLOQUE REPARADOR)
        def formatear_dias(dias):
            if pd.isna(dias):
                return "Nunca ingresó" # Reemplazamos el vacío por texto claro
            else:
                return str(int(dias)) # Cortamos los decimales y lo convertimos a texto

        # Armamos la tabla final lista para mostrar y exportar
        resultado_display = df[[col_nombre_completo, 'Estado', col_dias]].copy()
        resultado_display = resultado_display.sort_values(by=col_nombre_completo)
        
        # Aplicamos la limpieza de días a la columna final
        resultado_display[col_dias] = resultado_display[col_dias].apply(formatear_dias)
        
        # Renombramos columnas
        resultado_display.columns = ['Alumno', 'Clasificación', 'Días sin Ingresar']

        # Métricas
        st.markdown(f"#### Resumen de Asistencia (Límite: {dias_limite} días)")
        total_alumnos = len(resultado_display)
        total_presentes = len(resultado_display[resultado_display['Clasificación'] == "✅ PRESENTE"])
        total_ausentes = len(resultado_display[resultado_display['Clasificación'] == "❌ AUSENTE"])

        col1, col2, col3 = st.columns(3)
        col1.metric("Total Alumnos", total_alumnos)
        col2.metric("Presentes", total_presentes)
        col3.metric("Ausentes (En Riesgo)", total_ausentes, delta="Atención requerida", delta_color="inverse")

    # ==========================================
    # PESTAÑA 2: LISTADO DE ALUMNOS
    # ==========================================
    with tab2:
        st.markdown("#### Listado Detallado")
        
        def colorear_riesgo(val):
            if val == '❌ AUSENTE':
                return 'background-color: #ffcccc; color: #900000; font-weight: bold'
            return 'background-color: #e6ffe6; color: #006600'
            
        tabla_estilizada = resultado_display.style.map(colorear_riesgo, subset=['Clasificación'])
        st.dataframe(tabla_estilizada, use_container_width=True, hide_index=True)

    # ==========================================
    # PESTAÑA 3: CENTRO DE EXPORTACIÓN
    # ==========================================
    with tab3:
        st.markdown("#### Exportar Reporte")
        st.info("Los datos han sido formateados y están listos para descarga segura.")
        
        # Motor de Word optimizado (Ahora recibe texto limpio, no fallará)
        def crear_word(dataframe, titulo):
            doc = Document()
            doc.add_heading(f"Reporte de Asistencia: {titulo}", 0)
            
            tabla = doc.add_table(rows=1, cols=len(dataframe.columns))
            tabla.style = 'Table Grid'
            
            encabezados = tabla.rows[0].cells
            for i, columna in enumerate(dataframe.columns):
                encabezados[i].text = str(columna)
                
            for index, fila in dataframe.iterrows():
                celdas_fila = tabla.add_row().cells
                for i, valor in enumerate(fila):
                    # Forzamos todo a string por seguridad extrema
                    celdas_fila[i].text = str(valor)
                    
            archivo_memoria = io.BytesIO()
            doc.save(archivo_memoria)
            return archivo_memoria.getvalue()

        # Generamos los archivos en memoria
        word_listo = crear_word(resultado_display, file_name_no_ext)
        csv_listo = resultado_display.to_csv(index=False).encode('utf-8-sig')

        st.divider()
        col_btn1, col_btn2 = st.columns(2)
        
        with col_btn1:
            st.download_button(
                label="📥 Descargar Datos (CSV)",
                data=csv_listo,
                file_name=f"Asistencia_{file_name_no_ext}.csv",
                mime='text/csv',
            )
            
        with col_btn2:
            st.download_button(
                label="📥 Descargar Reporte Claro (Word)",
                data=word_listo,
                file_name=f"Reporte_{file_name_no_ext}.docx",
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )